import random
import docx
import re
import time


#产生随机数
def Number(m,n):
    return str(random.randint(m,n))



#产生随机符号
def caculate_symbol():
    symbol =['+','-','×','÷']
    return str(random.choice(symbol))

#生成题目列表
def Test(target_two,target_three):
    subject_all =[]
    for i in range(0,target_three):
        subject = ""
        for x in range(0,target_two):
            symbol = caculate_symbol()
            subject +=(Number(k,j) + symbol)

        subject = subject[0:-1] + "="
        subject_all.append(subject)

    return subject_all

#将答案保存在列表
def Answer(subject_all):
    answer = []
    for i in subject_all:
        a = ""
        c = re.sub('×',"*",i)
        b = re.sub('÷',"/",c)
        question = re.sub("=","",b)
        a += i + str("%.1f" %(eval(question)))
        answer.append(a)

    return answer



#保存为题目
def Save_Test(subject_all):
    doc = docx.Document()
    doc.add_heading("加减乘除速算(答案保留1位小数)")
    doc.add_paragraph('\n')
    for i in subject_all:
        doc.add_paragraph(i)

    doc.save("速算试卷.docx")
    print("\033[1;33;1m出题成功!\033[0m")
    #Python 输出颜色字符窜
    # \033[1;31;40m    <!--1-高亮显示 31-前景色红色  40-背景色黑色-->
    # 1m
    # \033[0m          <!--采用终端默认设置，即取消颜色设置-->
    '''
    前景色            背景色           颜色   显示方式           意义
-------------------------
0                终端默认设置
1                高亮显示
4                使用下划线
5                闪烁
7                反白显示
8                不可见
 
---------------------------------------
30                40              黑色
31                41              红色
32                42              绿色
33                43              黃色
34                44              蓝色
35                45              紫红色
36                46              青蓝色
37                47              白色
    '''


#保存答案
def Save_answer(answer):
    doc = docx.Document()
    doc.add_heading("加减乘除速算答案(答案保留一位小数)")
    for i in answer:
        doc.add_paragraph(i)

    doc.save("速算试卷答案.docx")
    print("\033[1;34;1m答案已给出!\033[0m","\033[1;31;1m题目为:%s个\033[0m" %len(answer))


#实现主要逻辑
def run():
    try:
        target_one = int(input("请运算数字的位数:"))
        global j
        j = 1
        for i in range(0,target_one):
            j *=10

        global k
        k = j // 10

        target_two = int(input("请输入运算数字的个数:"))
        target_three = int(input("请输入题数:"))

        subject_all = Test(target_two,target_three)

        Save_Test(subject_all)
        answer = Answer(subject_all)
        Save_answer(answer)

    except:
        print("\033[1;31;1m请按照提示正确输入!\033[0m")

if __name__ == '__main__':
    while True:
        run()
        time.sleep(0.05)








