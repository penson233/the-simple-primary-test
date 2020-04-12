import random
import docx
import re


#产生随机数
def Number(m,n):

    return str(random.randint(m,n))

#产生随机符号
def caculate_symbol():
    symbol =['+','-','×','÷']
    return str(random.choice(symbol))

#生成题目列表
def Test(target_two,target_three):
    subject_all =[]
    for i in range(0,target_three):
        subject = ""
        for x in range(0,target_two):
            symbol = caculate_symbol()
            subject +=(Number(k,j) + symbol)

        subject = subject[0:-1] + "="
        subject_all.append(subject)

    return subject_all

#将答案保存在列表
def Answer(subject_all):
    answer = []
    for i in subject_all:
        a = ""
        c = re.sub('×',"*",i)
        b = re.sub('÷',"/",c)
        question = re.sub("=","",b)
        a += i + str("%.2f" %(eval(question)))
        answer.append(a)

    return answer



#保存为题目
def Save_Test(subject_all):
    doc = docx.Document()
    doc.add_heading("加减乘除速算(答案保留2位小数)")
    for i in subject_all:
        doc.add_paragraph(i)

    doc.save("速算试卷.docx")

#保存答案
def Save_answer(answer):
    doc = docx.Document()
    doc.add_heading("加减乘除速算答案")
    for i in answer:
        doc.add_paragraph(i)

    doc.save("速算试卷答案.docx")

#实现主要逻辑
def run():
    target_one = int(input("你想要几位数:"))
    global j
    j = 1
    for i in range(0,target_one):
        j *=10

    global k
    k = j // 10

    target_two = int(input("你想要多少个数运算:"))
    target_three = int(input("你要出多少题:"))

    subject_all = Test(target_two,target_three)
    print(subject_all)

    Save_Test(subject_all)
    answer = Answer(subject_all)
    Save_answer(answer)


run()








